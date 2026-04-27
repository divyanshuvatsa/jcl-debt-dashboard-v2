"""
Document exports — Board Memo, per-lender one-pagers, compliance certificates,
CSV bundles.
"""

from __future__ import annotations
from datetime import date
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def _add_heading(doc, text, level=1, color=(255, 176, 0)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)


# =============================================================================
# BOARD MEMO
# =============================================================================
def build_board_memo(logic, ai_narrative: str = "") -> bytes:
    if Document is None:
        return b""

    doc = Document()

    # Title
    title = doc.add_heading("JCL DEBT PORTFOLIO — BOARD MEMORANDUM", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(255, 176, 0)

    p = doc.add_paragraph(f"Snapshot as of: {logic.as_of.strftime('%d-%b-%Y')}")
    p.runs[0].italic = True

    # Executive narrative
    _add_heading(doc, "Executive Summary", level=1)
    if ai_narrative:
        doc.add_paragraph(ai_narrative)
    else:
        from .ai_assistant import _fallback_board_narrative
        doc.add_paragraph(_fallback_board_narrative(logic))

    # KPIs
    _add_heading(doc, "Key Portfolio Metrics", level=1)
    health = logic.health_score()
    kpis = [
        ("Total Sanctioned",    f"₹{logic.total_sanctioned():.1f} Cr"),
        ("Effective Outstanding", f"₹{logic.total_outstanding():.1f} Cr"),
        ("Annual Cost",         f"₹{logic.annual_interest_commission():.2f} Cr"),
        ("Weighted Avg Cost",   f"{logic.wac_fb_plus_tl():.2f}%"),
        ("Health Score",        f"{health['composite']}/100"),
        ("WAM",                 f"{health['wam_months']:.1f} months"),
    ]
    table = doc.add_table(rows=len(kpis), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(kpis):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    # Lender breakdown
    _add_heading(doc, "Lender Breakdown", level=1)
    lenders = logic.lender_breakdown()
    table = doc.add_table(rows=len(lenders) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    hdrs = ["Lender", "Sanctioned (₹ Cr)", "% of Total", "Facilities"]
    for i, h in enumerate(hdrs):
        table.cell(0, i).text = h
    for i, l in enumerate(lenders, 1):
        table.cell(i, 0).text = l["lender"]
        table.cell(i, 1).text = f"{l['sanc']:.1f}"
        table.cell(i, 2).text = f"{l['pct_total']:.1f}%"
        table.cell(i, 3).text = str(l["facilities"])

    # Covenant status
    _add_heading(doc, "Covenant Status", level=1)
    cov_summary = logic.covenant_summary()
    doc.add_paragraph(
        f"Total: {cov_summary['total']} | "
        f"Compliant: {cov_summary['green']} | "
        f"Near Breach: {cov_summary['amber']} | "
        f"Breached: {cov_summary['red']} | "
        f"Compliance: {cov_summary['compliance_pct']:.1f}%"
    )

    near_breach = [c for c in logic.covenant_status() if "Near" in c["status"] or "Breached" in c["status"]]
    if near_breach:
        _add_heading(doc, "Watch Items", level=2)
        for c in near_breach:
            doc.add_paragraph(
                f"{c['lender']} — {c['covenant']}: actual {c['actual']:.2f} vs threshold {c['threshold']} "
                f"({c['headroom']:.1f}% headroom)",
                style="List Bullet"
            )

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Source: JCL Debt Model · For internal use only · Not financial advice"
    )
    foot.runs[0].italic = True
    foot.runs[0].font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# PER-LENDER ONE-PAGER
# =============================================================================
def build_lender_onepager(logic, lender: str) -> bytes:
    if Document is None:
        return b""

    doc = Document()

    title = doc.add_heading(f"{lender} — Facility One-Pager", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(255, 176, 0)

    doc.add_paragraph(f"As of: {logic.as_of.strftime('%d-%b-%Y')}").runs[0].italic = True

    # Facilities
    _add_heading(doc, "Facilities", level=1)
    facilities = [f for f in logic.facilities if f["lender"] == lender]
    table = doc.add_table(rows=len(facilities) + 1, cols=5)
    table.style = "Light Grid Accent 1"
    hdrs = ["Facility", "Type", "Sanc (₹ Cr)", "Rate", "Maturity"]
    for i, h in enumerate(hdrs):
        table.cell(0, i).text = h
    for i, f in enumerate(facilities, 1):
        table.cell(i, 0).text = f["facility"]
        table.cell(i, 1).text = f.get("category", "")
        table.cell(i, 2).text = f"{f.get('sanc_inr', 0):.1f}"
        rate = f.get("eff_rate")
        table.cell(i, 3).text = f"{rate:.2f}%" if rate else "TBD"
        mat = f.get("maturity")
        table.cell(i, 4).text = mat.strftime("%d-%b-%Y") if hasattr(mat, "strftime") else "—"

    # Covenants
    _add_heading(doc, "Covenants", level=1)
    lender_covs = [c for c in logic.covenant_status() if c["lender"] == lender]
    if lender_covs:
        table = doc.add_table(rows=len(lender_covs) + 1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Covenant", "Threshold", "Actual", "Status"]):
            table.cell(0, i).text = h
        for i, c in enumerate(lender_covs, 1):
            table.cell(i, 0).text = c["covenant"]
            table.cell(i, 1).text = f"{c['op']}{c['threshold']}"
            actual_str = f"{c['actual']:.2f}" if isinstance(c["actual"], (int, float)) else str(c["actual"])
            table.cell(i, 2).text = actual_str
            table.cell(i, 3).text = c["status"]

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# COMPLIANCE CERTIFICATE
# =============================================================================
def build_compliance_certificate(logic, lender: str) -> bytes:
    if Document is None:
        return b""

    doc = Document()
    title = doc.add_heading(f"COMPLIANCE CERTIFICATE — {lender}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(255, 176, 0)

    doc.add_paragraph(f"Date: {logic.as_of.strftime('%d-%b-%Y')}")
    doc.add_paragraph(f"To: {lender}")
    doc.add_paragraph(f"From: Jindal Coke Limited")
    doc.add_paragraph()

    doc.add_paragraph(
        "We hereby certify that, as of the above date, Jindal Coke Limited is in "
        "compliance with the financial covenants stipulated in the sanction letter "
        f"with {lender}, as detailed below:"
    )

    lender_covs = [c for c in logic.covenant_status() if c["lender"] == lender]
    if lender_covs:
        table = doc.add_table(rows=len(lender_covs) + 1, cols=4)
        table.style = "Medium Grid 1 Accent 1"
        for i, h in enumerate(["Covenant", "Threshold", "Actual", "Compliant?"]):
            table.cell(0, i).text = h
        for i, c in enumerate(lender_covs, 1):
            table.cell(i, 0).text = c["covenant"]
            table.cell(i, 1).text = f"{c['op']}{c['threshold']}"
            actual_str = f"{c['actual']:.2f}" if isinstance(c["actual"], (int, float)) else str(c["actual"])
            table.cell(i, 2).text = actual_str
            table.cell(i, 3).text = "Yes" if "Compliant" in c["status"] else "No"

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Authorised Signatory")
    doc.add_paragraph("Jindal Coke Limited")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
